"""
Extracts raw text from uploaded PDF, DOCX and TXT files.
Image-based files are routed to ocr_service instead.
"""
import os
from PyPDF2 import PdfReader
import docx
from app.utils.logger import logger
from app.utils.exceptions import ValidationError


def extract_text_from_pdf(file_path: str) -> str:
    try:
        reader = PdfReader(file_path)
        text_chunks = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text_chunks.append(page_text)
        return "\n".join(text_chunks).strip()
    except Exception as exc:  # noqa: BLE001
        logger.error(f"PDF extraction failed for {file_path}: {exc}")
        raise ValidationError("Could not extract text from PDF file") from exc


def extract_text_from_docx(file_path: str) -> str:
    try:
        document = docx.Document(file_path)
        paragraphs = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs).strip()
    except Exception as exc:  # noqa: BLE001
        logger.error(f"DOCX extraction failed for {file_path}: {exc}")
        raise ValidationError("Could not extract text from DOCX file") from exc


def extract_text_from_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception as exc:  # noqa: BLE001
        logger.error(f"TXT extraction failed for {file_path}: {exc}")
        raise ValidationError("Could not read text file") from exc


def extract_text(file_path: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    if ext == ".txt":
        return extract_text_from_txt(file_path)
    raise ValidationError(f"Unsupported file extension for text extraction: {ext}")
